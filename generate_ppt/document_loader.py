from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

import fitz
import markdown
from bs4 import BeautifulSoup
from docx import Document as DocxDocument


@dataclass
class SourceDocument:
    title: str
    text: str
    pages: list[str]
    source_name: str


def load_document(path: Path) -> SourceDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        pages = _load_pdf(path)
    elif suffix == ".docx":
        pages = _load_docx(path)
    elif suffix in {".md", ".markdown"}:
        pages = [_load_markdown(path)]
    elif suffix == ".txt":
        pages = [path.read_text(encoding="utf-8", errors="ignore")]
    else:
        raise ValueError(f"不支持的文件类型：{suffix}")

    text = _normalize("\n\n".join(pages))
    if not text.strip():
        raise ValueError("文档没有抽取到可用文本，可能是扫描版 PDF。")

    title = _guess_title(text, path.stem)
    return SourceDocument(title=title, text=text, pages=pages, source_name=path.name)


def _load_pdf(path: Path) -> list[str]:
    doc = fitz.open(path)
    return [_normalize(page.get_text("text")) for page in doc]


def _load_docx(path: Path) -> list[str]:
    doc = DocxDocument(path)
    lines = []
    for paragraph in doc.paragraphs:
        value = paragraph.text.strip()
        if value:
            lines.append(value)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return ["\n".join(lines)]


def _load_markdown(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    html = markdown.markdown(raw)
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text("\n")


def _guess_title(text: str, fallback: str) -> str:
    for line in text.splitlines():
        line = line.strip(" #\t")
        if 2 <= len(line) <= 40:
            return line
    return fallback


def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
